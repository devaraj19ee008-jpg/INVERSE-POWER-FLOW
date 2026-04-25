"""
DOCUMENT GENERATOR FOR POWER FLOW SYSTEM
Run this script to generate 3 Word documents (.docx)
Install requirement: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, size=11, space_after=6):
    """Add a paragraph with formatting"""
    p = doc.add_paragraph(text, style=style)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    return p

def create_calculation_document():
    """DOCUMENT 1: Main Calculations"""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('POWER FLOW CALCULATION METHODOLOGY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Mathematical Formulations for Hand Calculations & Automation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph(f'Document Version: 1.0 | Generated: {datetime.now().strftime("%Y-%m-%d")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_items = [
        '1. System Base Values and Per-Unit Conversion',
        '2. Admittance Matrix (Ybus) Construction',
        '   2.1 Transmission Line π-Model',
        '   2.2 Transformer Model',
        '   2.3 Complete Ybus Assembly Algorithm',
        '3. Power Injection Calculations',
        '   3.1 Complex Power from Voltage',
        '   3.2 Bus Power Balance Equations',
        '4. Branch Flow Calculations',
        '   4.1 Transmission Line Bidirectional Flows',
        '   4.2 Transformer Bidirectional Flows',
        '   4.3 Loss Calculations',
        '5. Loading and Rating Calculations',
        '6. Shunt Compensation Calculations',
        '7. Complete Calculation Workflow',
        '8. Hand Calculation Example',
        '9. Automation Code Template'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # SECTION 1: BASE VALUES
    doc.add_heading('1. SYSTEM BASE VALUES AND PER-UNIT CONVERSION', level=1)
    
    doc.add_heading('1.1 Base Values Definition', level=2)
    doc.add_paragraph('The per-unit system normalizes all electrical quantities:')
    
    add_code_block(doc, '''Base MVA (Sbase) = Given in input (typically 100 MVA)
Base Voltage (Vbase_kV) = Given per bus (e.g., 132 kV, 33 kV, 11 kV)
Base Current (Ibase) = Sbase / (√3 × Vbase_kV)
Base Impedance (Zbase) = Vbase_kV² / Sbase''')
    
    doc.add_heading('1.2 Per-Unit Conversion Formulas', level=2)
    
    # Table for conversions
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Quantity', 'Actual Value', 'Per-Unit Value']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '4472C4')
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = None
    
    conversions = [
        ('Voltage', 'V_kV', 'V_pu = V_kV / Vbase_kV'),
        ('Current', 'I_A', 'I_pu = I_A / Ibase'),
        ('Impedance', 'Z_Ω', 'Z_pu = Z_Ω / Zbase'),
        ('Power', 'S_MVA', 'S_pu = S_MVA / Sbase'),
        ('Reactive Power', 'Q_Mvar', 'Q_pu = Q_Mvar / Sbase')
    ]
    for i, (qty, actual, pu) in enumerate(conversions, 1):
        table.rows[i].cells[0].text = qty
        table.rows[i].cells[1].text = actual
        table.rows[i].cells[2].text = pu
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 Input Data Per-Unit Values', level=2)
    doc.add_paragraph('The input file provides impedance values already in per-unit on system base:')
    add_code_block(doc, '''LINE_DATA format: [line_num, name, from_bus, to_bus, length_km, 
                 R_per_km, X_per_km, B_per_km, rateA_MVA, status, area, zone, owner]

Total Line Impedance (pu):
    R_total_pu = R_per_km × length_km
    X_total_pu = X_per_km × length_km  
    B_total_pu = B_per_km × length_km

TRANSFORMER_DATA format: [xfmr_num, name, from_bus, to_bus, R_pu, X_pu, 
                          tap_ratio, rateA_MVA, phase_shift_deg, ...]

Note: Transformer R and X are already total values in pu''')
    
    doc.add_page_break()
    
    # SECTION 2: YBUS CONSTRUCTION
    doc.add_heading('2. ADMITTANCE MATRIX (Ybus) CONSTRUCTION', level=1)
    
    doc.add_paragraph('The Ybus matrix is the core of power flow calculations. It relates bus voltages to injected currents:')
    add_code_block(doc, '''[I_bus] = [Ybus] × [V_bus]

Where:
    I_bus = Vector of bus injection currents (n×1 complex)
    V_bus = Vector of bus voltages (n×1 complex)
    Ybus  = Admittance matrix (n×n complex)
    n     = Number of buses''')
    
    doc.add_heading('2.1 Transmission Line π-Model', level=2)
    
    doc.add_paragraph('Physical Model Diagram:')
    add_code_block(doc, '''        Bus i                              Bus j
           │                                 │
           ├──┬────── R+jX ──────┬───────────┤
           │  │                  │           │
          jB/2│                 jB/2         │
           │  │                  │           │
           └──┴──────────────────┴───────────┘
           
Where:
    R = Series resistance (pu)
    X = Series reactance (pu)
    B = Total line charging susceptance (pu)''')
    
    doc.add_paragraph('Mathematical Derivation:')
    add_code_block(doc, '''Step 1: Calculate series admittance
    Z_series = R + jX  (complex impedance)
    y_series = 1 / Z_series = 1 / (R + jX)
    
    Rationalize the denominator:
    y_series = (R - jX) / (R² + X²)
    
    Let D = R² + X²
    G_series = R / D      (series conductance)
    B_series = -X / D     (series susceptance)
    
    y_series = G_series + jB_series

Step 2: Shunt admittance (per side)
    y_shunt = j(B/2)      (purely reactive)

Step 3: Ybus contributions from this line
    Ybus[i][i] += y_series + j(B/2)   (diagonal element at bus i)
    Ybus[j][j] += y_series + j(B/2)   (diagonal element at bus j)
    Ybus[i][j] -= y_series             (off-diagonal element)
    Ybus[j][i] -= y_series             (off-diagonal element, symmetric)''')
    
    doc.add_heading('2.2 Transformer Model', level=2)
    
    doc.add_paragraph('The transformer is modeled as an ideal transformer with series impedance:')
    add_code_block(doc, '''        Bus i (Primary)              Bus j (Secondary)
           │                              │
           │    ┌──────────────┐          │
           ├───┤  Tap Ratio a  ├──────────┤
           │    └──────┬───────┘          │
           │           │                  │
           │         R+jX                 │
           │           │                  │
           └───────────┴──────────────────┘
           
Where:
    a = tap_ratio (typically 0.9 to 1.1)
    R = Transformer resistance (pu)
    X = Transformer reactance (pu)''')
    
    doc.add_paragraph('Mathematical Derivation:')
    add_code_block(doc, '''Step 1: Calculate series admittance (same as line)
    Z_xfmr = R + jX
    y_series = 1 / (R + jX) = (R - jX) / (R² + X²)

Step 2: Ybus contributions with tap ratio
    Ybus[i][i] += y_series / a²         (primary side, adjusted by tap)
    Ybus[j][j] += y_series               (secondary side, no adjustment)
    Ybus[i][j] -= y_series / a           (off-diagonal, divided by tap)
    Ybus[j][i] -= y_series / a           (symmetric off-diagonal)

Key Insight: The tap ratio a affects only the primary side diagonal 
element and both off-diagonal elements. This asymmetry is what 
allows transformers to control voltage.''')
    
    doc.add_heading('2.3 Complete Ybus Assembly Algorithm', level=2)
    
    add_code_block(doc, '''ALGORITHM: Build Ybus Matrix
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

INPUT:
    - n_buses: Number of buses
    - branches: List of {from_bus, to_bus, r, x, b, ratio}
    - bus_index_map: Dictionary mapping bus_num → matrix index

OUTPUT:
    - Ybus: n×n complex matrix

STEPS:
1. Initialize Ybus as n×n zero matrix (complex)
   
2. FOR EACH branch in branches:
   a. Get matrix indices:
      i = bus_index_map[from_bus]
      j = bus_index_map[to_bus]
   
   b. Calculate series admittance:
      Z = complex(r, x)
      IF Z ≠ 0:
         y = 1.0 / Z
      ELSE:
         SKIP this branch (invalid)
   
   c. IF branch.ratio == 0 (TRANSMISSION LINE):
      Ybus[i][i] += y + complex(0, b)
      Ybus[j][j] += y + complex(0, b)
      Ybus[i][j] -= y
      Ybus[j][i] -= y
   
   d. ELSE (TRANSFORMER):
      a = branch.ratio
      Ybus[i][i] += y / (a²)
      Ybus[j][j] += y
      Ybus[i][j] -= y / a
      Ybus[j][i] -= y / a

3. RETURN Ybus

Python Implementation:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
import numpy as np

def build_ybus(n_buses, branches, bus_index_map):
    Ybus = np.zeros((n_buses, n_buses), dtype=complex)
    
    for branch in branches:
        i = bus_index_map[branch['from_bus']]
        j = bus_index_map[branch['to_bus']]
        r, x, b = branch['r'], branch['x'], branch['b']
        ratio = branch['ratio']
        
        z = complex(r, x)
        if z != 0:
            y = 1.0 / z
            
            if ratio == 0:  # Line
                Ybus[i][i] += y + complex(0, b)
                Ybus[j][j] += y + complex(0, b)
                Ybus[i][j] -= y
                Ybus[j][i] -= y
            else:  # Transformer
                Ybus[i][i] += y / (ratio ** 2)
                Ybus[j][j] += y
                Ybus[i][j] -= y / ratio
                Ybus[j][i] -= y / ratio
    
    return Ybus''')
    
    doc.add_page_break()
    
    # SECTION 3: POWER INJECTION
    doc.add_heading('3. POWER INJECTION CALCULATIONS', level=1)
    
    doc.add_heading('3.1 Complex Power from Voltage and Current', level=2)
    
    add_code_block(doc, '''FUNDAMENTAL EQUATION:
    S = V × I*

Where:
    S = Complex power (VA)
    V = Complex voltage (V)
    I* = Complex conjugate of current (A)

In Per-Unit:
    S_pu = V_pu × I_pu*

From Ybus:
    I = Ybus × V

Therefore:
    S = V × (Ybus × V)* = V × Ybus* × V*

For bus k:
    S_k = V_k × Σ(j=1 to n) [Y_kj* × V_j*]

Expanding into real and imaginary:
    S_k = P_k + jQ_k

Where:
    P_k = Real part of S_k (active power)
    Q_k = Imaginary part of S_k (reactive power)''')
    
    doc.add_heading('3.2 Detailed Power Equations', level=2)
    
    add_code_block(doc, '''For bus k with voltage V_k = |V_k|∠δ_k:

P_k = |V_k| × Σ(j=1 to n) |V_j| × |Y_kj| × cos(δ_k - δ_j - θ_kj)

Q_k = |V_k| × Σ(j=1 to n) |V_j| × |Y_kj| × sin(δ_k - δ_j - θ_kj)

Where:
    |V_k| = Voltage magnitude at bus k (pu)
    δ_k   = Voltage angle at bus k (radians)
    |Y_kj| = Magnitude of admittance element Y_kj
    θ_kj   = Angle of admittance element Y_kj

Alternative Matrix Form (used in code):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Given:
    V_complex = V_mag × exp(j × V_angle_rad)

Calculate:
    I_injection = Ybus @ V_complex      (matrix multiplication)
    S_complex = V_complex × conj(I_injection)
    
    P = real(S_complex)                 (active power, pu)
    Q = imag(S_complex)                 (reactive power, pu)
    
Convert to MW/Mvar:
    P_MW = P_pu × Sbase_MVA
    Q_Mvar = Q_pu × Sbase_MVA

SIGN CONVENTION:
    Positive P, Q → Power flowing INTO the bus (generation)
    Negative P, Q → Power flowing OUT of bus (load)''')
    
    doc.add_heading('3.3 Python Implementation', level=2)
    
    add_code_block(doc, '''def calculate_power_injection(Ybus, V_mag, V_angle_deg, baseMVA):
    """
    Calculate power injection at all buses
    
    Parameters:
        Ybus: n×n complex admittance matrix
        V_mag: array of voltage magnitudes (pu)
        V_angle_deg: array of voltage angles (degrees)
        baseMVA: system base MVA
    
    Returns:
        P_MW: array of active power injections (MW)
        Q_Mvar: array of reactive power injections (Mvar)
    """
    # Convert angle to radians
    V_angle_rad = np.radians(V_angle_deg)
    
    # Form complex voltage vector
    V_complex = V_mag * np.exp(1j * V_angle_rad)
    
    # Calculate injected current
    I_inj = Ybus @ V_complex
    
    # Calculate complex power: S = V × I*
    S = V_complex * np.conj(I_inj)
    
    # Extract real and imaginary parts
    P_pu = S.real
    Q_pu = S.imag
    
    # Convert to MW and Mvar
    P_MW = P_pu * baseMVA
    Q_Mvar = Q_pu * baseMVA
    
    return P_MW, Q_Mvar''')
    
    doc.add_page_break()
    
    # SECTION 4: BRANCH FLOWS
    doc.add_heading('4. BRANCH FLOW CALCULATIONS', level=1)
    
    doc.add_heading('4.1 Transmission Line Bidirectional Flows', level=2)
    
    add_code_block(doc, '''PHYSICAL MODEL (π-model):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        Bus i                              Bus j
           │                                 │
           ├──┬────── y ──────┬─────────────┤
           │  │               │             │
         jB/2│              jB/2           │
           │  │               │             │
           └──┴───────────────┴─────────────┘

Where:
    y = 1/(R+jX) = series admittance
    B = total line charging susceptance

FORWARD FLOW (i → j):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Current through series element (i to j):
    I_series_ij = (V_i - V_j) × y

Shunt charging current at bus i:
    I_shunt_i = V_i × j(B/2)

Total current leaving bus i:
    I_from_i = I_series_ij + I_shunt_i
    I_from_i = (V_i - V_j) × y + V_i × j(B/2)

Complex power from bus i:
    S_ij = V_i × I_from_i*
    S_ij = V_i × [(V_i - V_j) × y + V_i × j(B/2)]*

REVERSE FLOW (j → i):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Current through series element (j to i):
    I_series_ji = (V_j - V_i) × y

Shunt charging current at bus j:
    I_shunt_j = V_j × j(B/2)

Total current leaving bus j:
    I_from_j = I_series_ji + I_shunt_j

Complex power from bus j:
    S_ji = V_j × I_from_j*

POWER LOSSES:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Total losses = S_ij + S_ji
(Note: This is positive for real power loss)

Loss_MW = real(S_ij + S_ji) × baseMVA
Loss_Mvar = imag(S_ij + S_ji) × baseMVA''')
    
    doc.add_heading('4.2 Transformer Bidirectional Flows', level=2)
    
    add_code_block(doc, '''TRANSFORMER MODEL:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        Bus i (Primary, tap side)      Bus j (Secondary)
           │                              │
           │    ┌──────────────┐          │
           ├───┤  Tap Ratio a  ├──────────┤
           │    └──────┬───────┘          │
           │           │                  │
           │           y                  │
           │           │                  │
           └───────────┴──────────────────┘

Where:
    a = tap_ratio
    y = 1/(R+jX) = series admittance

FORWARD FLOW (i → j):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Current through transformer (accounting for tap):
    I_fwd = (V_i/a - V_j) × y

Complex power leaving bus i:
    S_fwd = V_i × I_fwd*
    S_fwd = V_i × [(V_i/a - V_j) × y]*

REVERSE FLOW (j → i):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Current from secondary side (accounting for tap):
    I_rev = (V_j × a - V_i) × (y / a²)

Complex power leaving bus j:
    S_rev = V_j × I_rev*
    S_rev = V_j × [(V_j × a - V_i) × (y / a²)]*

POWER LOSSES:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Loss = S_fwd + S_rev
Loss_MW = real(S_fwd + S_rev) × baseMVA

IMPORTANT: The asymmetry in forward/reverse equations is due 
to the tap ratio being on the primary side.''')
    
    doc.add_heading('4.3 Complete Python Implementation', level=2)
    
    add_code_block(doc, '''def calculate_branch_flows(Ybus, branches, bus_index_map, 
                           V_mag, V_angle_deg, baseMVA):
    """
    Calculate bidirectional power flows on all branches
    
    Returns: List of flow dictionaries for each branch
    """
    V_angle_rad = np.radians(V_angle_deg)
    V_complex = V_mag * np.exp(1j * V_angle_rad)
    
    flows = []
    
    for branch in branches:
        i = bus_index_map[branch['from_bus']]
        j = bus_index_map[branch['to_bus']]
        V_i = V_complex[i]
        V_j = V_complex[j]
        
        r, x, b = branch['r'], branch['x'], branch['b']
        ratio = branch['ratio']
        rateA = branch.get('rateA', 9999)
        
        z = complex(r, x)
        y = 1.0 / z if z != 0 else 0
        
        if ratio == 0:  # TRANSMISSION LINE
            # Forward flow (from_bus → to_bus)
            I_series = (V_i - V_j) * y
            I_shunt = V_i * complex(0, b/2)
            S_fwd = V_i * np.conj(I_series + I_shunt)
            
            # Reverse flow (to_bus → from_bus)
            I_series_rev = (V_j - V_i) * y
            I_shunt_rev = V_j * complex(0, b/2)
            S_rev = V_j * np.conj(I_series_rev + I_shunt_rev)
            
        else:  # TRANSFORMER
            # Forward flow
            I_fwd = (V_i / ratio - V_j) * y
            S_fwd = V_i * np.conj(I_fwd)
            
            # Reverse flow
            I_rev = (V_j * ratio - V_i) * (y / ratio**2)
            S_rev = V_j * np.conj(I_rev)
        
        losses = S_fwd + S_rev
        MVA_fwd = abs(S_fwd) * baseMVA
        loading = (MVA_fwd / rateA * 100) if rateA > 0 else 0
        
        flows.append({
            'from_bus': branch['from_bus'],
            'to_bus': branch['to_bus'],
            'P_fwd_MW': S_fwd.real * baseMVA,
            'Q_fwd_Mvar': S_fwd.imag * baseMVA,
            'P_rev_MW': S_rev.real * baseMVA,
            'Q_rev_Mvar': S_rev.imag * baseMVA,
            'losses_MW': losses.real * baseMVA,
            'MVA_fwd': MVA_fwd,
            'loading_pct': loading,
            'tap_ratio': ratio if ratio != 0 else None
        })
    
    return flows''')
    
    doc.add_page_break()
    
    # SECTION 5: LOADING CALCULATIONS
    doc.add_heading('5. LOADING AND RATING CALCULATIONS', level=1)
    
    add_code_block(doc, '''LINE LOADING PERCENTAGE:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Step 1: Calculate apparent power (MVA)
    P_fwd = |P_fwd_MW|           (absolute value)
    Q_fwd = |Q_fwd_Mvar|         (absolute value)
    MVA_flow = √(P_fwd² + Q_fwd²)

Step 2: Get rating
    rateA = Line rating from input (MVA)
    If rateA = 0, set to 9999 (no limit)

Step 3: Calculate loading percentage
    Loading_% = (MVA_flow / rateA) × 100

Step 4: Status determination
    IF Loading_% > 100:    Status = "OVERLOADED"
    ELIF Loading_% > 80:   Status = "HIGH"
    ELIF Loading_% > 60:   Status = "MODERATE"
    ELSE:                  Status = "NORMAL"

TRANSFORMER LOADING:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Same formula applies, but use transformer rateA:
    Loading_% = (MVA_fwd / Xfmr_rateA) × 100

TAP POSITION CALCULATION:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Given:
    tap_ratio = Current tap ratio (e.g., 1.025)
    min_tap = Minimum tap (e.g., 0.9)
    step_size = Tap step size (e.g., 0.0125)

Calculate:
    tap_position = (tap_ratio - min_tap) / step_size
    tap_step = int(tap_position)

Example:
    tap_ratio = 1.025, min_tap = 0.9, step_size = 0.0125
    tap_step = (1.025 - 0.9) / 0.0125 = 10
    → Transformer is at tap position 10''')
    
    doc.add_page_break()
    
    # SECTION 6: SHUNT COMPENSATION
    doc.add_heading('6. SHUNT COMPENSATION CALCULATIONS', level=1)
    
    add_code_block(doc, '''CAPACITOR BANK OUTPUT:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Given:
    Q_rated = Rated reactive power at nominal voltage (Mvar)
    V_actual = Actual bus voltage (pu)

Actual output (varies with V²):
    Q_actual = Q_rated × V_actual²

Reason: Capacitor impedance Z_c = 1/(jωC) is constant
        Current I = V/Z_c = V × jωC
        Power Q = V × I = V² × ωC ∝ V²

REACTOR OUTPUT:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Same formula applies (reactor absorbs reactive power):
    Q_absorbed = Q_rated × V_actual²

CONVERSION FROM BUS SHUNT DATA:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
In PYPOWER format, bus shunt is in pu on system base:
    B_shunt_pu = Value from bus data column 5

If B_shunt_pu < 0:  → Capacitor (negative susceptance = capacitive)
    Q_cap_Mvar = |B_shunt_pu| × baseMVA

If B_shunt_pu > 0:  → Reactor (positive susceptance = inductive)
    Q_react_Mvar = |B_shunt_pu| × baseMVA''')
    
    doc.add_page_break()
    
    # SECTION 7: COMPLETE WORKFLOW
    doc.add_heading('7. COMPLETE CALCULATION WORKFLOW', level=1)
    
    add_code_block(doc, '''STEP-BY-STEP CALCULATION PROCEDURE:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

┌─────────────────────────────────────────────────────────┐
│ STEP 1: DATA PREPARATION                                │
├─────────────────────────────────────────────────────────┤
│ • Read input file (BUS_DATA, LINE_DATA, etc.)          │
│ • Extract baseMVA (typically 100)                       │
│ • Build bus_index_map: {bus_num → 0,1,2,...}          │
│ • Calculate line totals: R×L, X×L, B×L                 │
└─────────────────────────────────────────────────────────┘
                           ↓
┌─────────────────────────────────────────────────────────┐
│ STEP 2: BUILD YBUS MATRIX                              │
├─────────────────────────────────────────────────────────┤
│ • Initialize n×n zero matrix (complex)                 │
│ • For each branch:                                      │
│   - Calculate y = 1/(R+jX)                             │
│   - If line: Add y±jB to diagonals, -y to off-diag    │
│   - If transformer: Adjust by tap ratio a              │
└─────────────────────────────────────────────────────────┘
                           ↓
┌─────────────────────────────────────────────────────────┐
│ STEP 3: FORM VOLTAGE VECTOR                            │
├─────────────────────────────────────────────────────────┤
│ • V_mag[i] = Voltage magnitude at bus i (pu)           │
│ • V_angle[i] = Voltage angle at bus i (degrees)        │
│ • V_complex = V_mag × exp(j × radians(V_angle))        │
└─────────────────────────────────────────────────────────┘
                           ↓
┌─────────────────────────────────────────────────────────┐
│ STEP 4: CALCULATE BUS INJECTIONS                       │
├─────────────────────────────────────────────────────────┤
│ • I_inj = Ybus × V_complex                             │
│ • S = V_complex × conj(I_inj)                          │
│ • P_MW = real(S) × baseMVA                             │
│ • Q_Mvar = imag(S) × baseMVA                           │
└─────────────────────────────────────────────────────────┘
                           ↓
┌─────────────────────────────────────────────────────────┐
│ STEP 5: CALCULATE BRANCH FLOWS                         │
├─────────────────────────────────────────────────────────┤
│ • For each line:                                        │
│   - Forward: S_fwd = Vi × [(Vi-Vj)×y + Vi×jB/2]*      │
│   - Reverse: S_rev = Vj × [(Vj-Vi)×y + Vj×jB/2]*      │
│ • For each transformer:                                 │
│   - Forward: S_fwd = Vi × [(Vi/a - Vj)×y]*             │
│   - Reverse: S_rev = Vj × [(Vj×a - Vi)×(y/a²)]*       │
│ • Losses = S_fwd + S_rev                               │
│ • Loading = |S_fwd| × baseMVA / rateA × 100            │
└─────────────────────────────────────────────────────────┘
                           ↓
┌─────────────────────────────────────────────────────────┐
│ STEP 6: PROCESS SHUNTS                                 │
├─────────────────────────────────────────────────────────┤
│ • Capacitor: Q_out = Q_rated × V²                      │
│ • Reactor: Q_abs = Q_rated × V²                        │
└─────────────────────────────────────────────────────────┘
                           ↓
┌─────────────────────────────────────────────────────────┐
│ STEP 7: VALIDATE & OUTPUT                              │
├─────────────────────────────────────────────────────────┤
│ • Check for NaN/Inf values                              │
│ • Check voltage limits (0.85 - 1.15 pu)                │
│ • Generate output files (CSV, PY, TXT)                 │
└─────────────────────────────────────────────────────────┘''')
    
    doc.add_page_break()
    
    # SECTION 8: HAND CALCULATION EXAMPLE
    doc.add_heading('8. HAND CALCULATION EXAMPLE', level=1)
    
    doc.add_paragraph('Example: 3-Bus System')
    add_code_block(doc, '''SYSTEM DATA:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Base MVA = 100

Bus Data:
  Bus 1: V = 1.05∠0° pu (Slack), base_kV = 132
  Bus 2: V = 1.00∠?° pu (PQ), base_kV = 132
  Bus 3: V = 1.00∠?° pu (PQ), base_kV = 33

Line Data:
  Line 1-2: R=0.02, X=0.06, B=0.06 pu
  Line 1-3: R=0.08, X=0.24, B=0.05 pu (via transformer)
  Line 2-3: R=0.06, X=0.18, B=0.04 pu

Transformer (Line 1-3 has tap ratio):
  Tap ratio = 0.95 (step-down from 132kV to 33kV)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 1: Calculate series admittances
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Line 1-2: Z = 0.02 + j0.06
  D = 0.02² + 0.06² = 0.004
  y₁₂ = (0.02 - j0.06) / 0.004 = 5 - j15

Line 1-3: Z = 0.08 + j0.24
  D = 0.08² + 0.24² = 0.064
  y₁₃ = (0.08 - j0.24) / 0.064 = 1.25 - j3.75

Line 2-3: Z = 0.06 + j0.18
  D = 0.06² + 0.18² = 0.036
  y₂₃ = (0.06 - j0.18) / 0.036 = 1.667 - j5.0

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 2: Build Ybus (3×3 matrix)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Y₁₁ = y₁₂ + j(0.06/2) + y₁₃/(0.95²) + j(0.05/2)
     = (5-j15) + j0.03 + (1.25-j3.75)/0.9025 + j0.025
     = 5 + 1.385 + j(-15 + 0.03 - 4.155 + 0.025)
     = 6.385 - j19.1

Y₂₂ = y₁₂ + j(0.06/2) + y₂₃ + j(0.04/2)
     = (5-j15) + j0.03 + (1.667-j5.0) + j0.02
     = 6.667 - j19.95

Y₃₃ = y₁₃ + j(0.05/2) + y₂₃ + j(0.04/2)
     = (1.25-j3.75) + j0.025 + (1.667-j5.0) + j0.02
     = 2.917 - j8.705

Y₁₂ = Y₂₁ = -y₁₂ = -5 + j15

Y₁₃ = -y₁₃/0.95 = -(1.25-j3.75)/0.95 = -1.316 + j3.947

Y₂₃ = Y₃₂ = -y₂₃ = -1.667 + j5.0

Ybus = | 6.385-j19.1   -5+j15       -1.316+j3.947 |
       | -5+j15         6.667-j19.95 -1.667+j5.0   |
       | -1.316+j3.947  -1.667+j5.0   2.917-j8.705 |

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 3: Calculate power injection at Bus 1 (given V)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Assume V₂ = 1.00∠-2° pu, V₃ = 0.98∠-5° pu (from power flow)

V = [1.05∠0°, 1.00∠-2°, 0.98∠-5°]

Convert to rectangular:
V₁ = 1.05 + j0
V₂ = 0.9994 - j0.0349
V₃ = 0.9763 - j0.0854

I₁ = Y₁₁×V₁ + Y₁₂×V₂ + Y₁₃×V₃
   = (6.385-j19.1)(1.05) + (-5+j15)(0.9994-j0.0349) + ...
   
S₁ = V₁ × I₁* = P₁ + jQ₁

(Continue calculation for all buses...)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 4: Calculate line flow 1→2
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

I_series = (V₁ - V₂) × y₁₂ = (1.05+j0 - 0.9994-j0.0349) × (5-j15)
         = (0.0506 + j0.0349) × (5-j15)

I_shunt = V₁ × j(B/2) = (1.05) × j(0.03) = j0.0315

S₁₂ = V₁ × (I_series + I_shunt)*

This gives P_fwd and Q_fwd in pu. Multiply by 100 for MW/Mvar.''')
    
    doc.add_page_break()
    
    # SECTION 9: AUTOMATION TEMPLATE
    doc.add_heading('9. COMPLETE AUTOMATION CODE TEMPLATE', level=1)
    
    add_code_block(doc, '''"""
Standalone Power Flow Calculator
Copy this code to create your own power flow tool
"""

import numpy as np

class PowerFlowCalculator:
    def __init__(self, baseMVA=100.0):
        self.baseMVA = baseMVA
        self.Ybus = None
        self.bus_index_map = {}
        self.n_buses = 0
        
    def load_system(self, buses, branches):
        """
        buses: dict {bus_num: {'V_init': float, 'angle_init': float, 'base_kV': float}}
        branches: list of dict {from_bus, to_bus, r, x, b, ratio, rateA}
        """
        self.buses = buses
        self.branches = branches
        self.n_buses = len(buses)
        self.bus_index_map = {b: i for i, b in enumerate(sorted(buses.keys()))}
        self._build_ybus()
        
    def _build_ybus(self):
        self.Ybus = np.zeros((self.n_buses, self.n_buses), dtype=complex)
        for br in self.branches:
            i = self.bus_index_map[br['from_bus']]
            j = self.bus_index_map[br['to_bus']]
            z = complex(br['r'], br['x'])
            if z == 0: continue
            y = 1.0 / z
            if br['ratio'] == 0:  # Line
                self.Ybus[i,i] += y + complex(0, br['b'])
                self.Ybus[j,j] += y + complex(0, br['b'])
                self.Ybus[i,j] -= y
                self.Ybus[j,i] -= y
            else:  # Transformer
                a = br['ratio']
                self.Ybus[i,i] += y / (a**2)
                self.Ybus[j,j] += y
                self.Ybus[i,j] -= y / a
                self.Ybus[j,i] -= y / a
    
    def solve(self, V_mag=None, V_angle=None):
        """Solve power flow for given voltages"""
        if V_mag is None:
            V_mag = np.array([self.buses[b]['V_init'] 
                             for b in sorted(self.buses.keys())])
        if V_angle is None:
            V_angle = np.array([self.buses[b]['angle_init'] 
                               for b in sorted(self.buses.keys())])
        
        V_rad = np.radians(V_angle)
        V = V_mag * np.exp(1j * V_rad)
        I = self.Ybus @ V
        S = V * np.conj(I)
        
        results = {
            'P_MW': S.real * self.baseMVA,
            'Q_Mvar': S.imag * self.baseMVA,
            'V_pu': V_mag,
            'V_angle_deg': V_angle,
            'branches': self._calc_flows(V)
        }
        return results
    
    def _calc_flows(self, V):
        flows = []
        for br in self.branches:
            i = self.bus_index_map[br['from_bus']]
            j = self.bus_index_map[br['to_bus']]
            Vi, Vj = V[i], V[j]
            y = 1.0 / complex(br['r'], br['x'])
            
            if br['ratio'] == 0:  # Line
                Ish = Vi * complex(0, br['b']/2)
                S_fwd = Vi * np.conj((Vi-Vj)*y + Ish)
                Ish_j = Vj * complex(0, br['b']/2)
                S_rev = Vj * np.conj((Vj-Vi)*y + Ish_j)
            else:  # Transformer
                a = br['ratio']
                S_fwd = Vi * np.conj((Vi/a - Vj)*y)
                S_rev = Vj * np.conj((Vj*a - Vi)*(y/a**2))
            
            MVA = abs(S_fwd) * self.baseMVA
            rateA = br.get('rateA', 9999)
            
            flows.append({
                'from': br['from_bus'], 'to': br['to_bus'],
                'P_fwd_MW': S_fwd.real * self.baseMVA,
                'Q_fwd_Mvar': S_fwd.imag * self.baseMVA,
                'P_rev_MW': S_rev.real * self.baseMVA,
                'Q_rev_Mvar': S_rev.imag * self.baseMVA,
                'loss_MW': (S_fwd + S_rev).real * self.baseMVA,
                'MVA': MVA,
                'loading_%': MVA/rateA*100 if rateA > 0 else 0
            })
        return flows


# USAGE EXAMPLE
if __name__ == "__main__":
    pf = PowerFlowCalculator(baseMVA=100)
    
    # Define system
    buses = {
        1: {'V_init': 1.05, 'angle_init': 0, 'base_kV': 132},
        2: {'V_init': 1.00, 'angle_init': -2, 'base_kV': 132},
        3: {'V_init': 0.98, 'angle_init': -5, 'base_kV': 33}
    }
    
    branches = [
        {'from_bus': 1, 'to_bus': 2, 'r': 0.02, 'x': 0.06, 'b': 0.06, 'ratio': 0, 'rateA': 100},
        {'from_bus': 1, 'to_bus': 3, 'r': 0.08, 'x': 0.24, 'b': 0.05, 'ratio': 0.95, 'rateA': 50},
        {'from_bus': 2, 'to_bus': 3, 'r': 0.06, 'x': 0.18, 'b': 0.04, 'ratio': 0, 'rateA': 80}
    ]
    
    pf.load_system(buses, branches)
    results = pf.solve()
    
    print("Bus Injections:")
    for i, bus in enumerate(sorted(buses.keys())):
        print(f"  Bus {bus}: P={results['P_MW'][i]:.2f} MW, Q={results['Q_Mvar'][i]:.2f} Mvar")
    
    print("\\nBranch Flows:")
    for f in results['branches']:
        print(f"  {f['from']}→{f['to']}: P={f['P_fwd_MW']:.2f} MW, Loading={f['loading_%']:.1f}%")
''')
    
    # Save document
    doc.save('01_Main_Calculations.docx')
    print("Created: 01_Main_Calculations.docx")


def create_user_manual_document():
    """DOCUMENT 2: User Manual"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('INVERSE POWER FLOW SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('USER MANUAL', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Version 1.0 | {datetime.now().strftime("%Y-%m-%d")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # SECTION 1: INTRODUCTION
    doc.add_heading('1. INTRODUCTION', level=1)
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph('The Inverse Power Flow System is a modular Python-based tool for power system analysis. Unlike traditional power flow solvers that iterate to find voltages from given powers, this system:')
    p = doc.add_paragraph()
    p.add_run('• Takes voltage profiles as input').bold = False
    p = doc.add_paragraph()
    p.add_run('• Calculates the resulting power flows and injections').bold = False
    p = doc.add_paragraph()
    p.add_run('• Generates multiple random samples for statistical analysis').bold = False
    p = doc.add_paragraph()
    p.add_run('• Produces comprehensive reports in multiple formats').bold = False
    
    doc.add_heading('1.2 System Components', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['File', 'Function', 'Modifiable?']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '2E75B6')
    
    components = [
        ('frompypowerformatetothis.py', 'Converts PYPOWER files to custom format', 'Yes'),
        ('input_reader.py', 'Reads and parses input files', 'Yes'),
        ('power_flow_engine.py', 'Core calculation engine', 'NO'),
        ('report_writer.py', 'Generates output reports', 'Yes'),
        ('run_power_flow.py', 'Main execution script', 'Yes')
    ]
    for i, (f, func, mod) in enumerate(components, 1):
        table.rows[i].cells[0].text = f
        table.rows[i].cells[1].text = func
        table.rows[i].cells[2].text = mod
    
    doc.add_page_break()
    
    # SECTION 2: INPUTS
    doc.add_heading('2. INPUT SPECIFICATIONS', level=1)
    
    doc.add_heading('2.1 Input File Format', level=2)
    doc.add_paragraph('The system accepts Python (.py) files with specific variable names:')
    
    add_code_block(doc, '''Required Variables:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BASE_MVA = 100                    # System base MVA

BUS_DATA = [                   # Bus information
    [bus_num, name, type, base_kV, V_init, angle, G, B, area, zone, owner],
    ...
]

GENERATOR_DATA = [              # Generator information
    [gen_num, name, type, bus, P_out, Q_out, V_set, Qmin, Qmax, status, area, zone],
    ...
]

LOAD_DATA = [                   # Load information
    [load_num, name, bus, P_demand, Q_demand, model, area, zone],
    ...
]

LINE_DATA = [                   # Transmission line information
    [line_num, name, from_bus, to_bus, length_km, R_per_km, X_per_km, 
     B_per_km, rateA_MVA, status, area, zone, owner],
    ...
]

TRANSFORMER_DATA = [            # Transformer information
    [xfmr_num, name, from_bus, to_bus, R_pu, X_pu, tap_ratio, rateA_MVA,
     phase_shift, min_tap, max_tap, step_size, status, area, zone, owner],
    ...
]

# Optional (can be empty lists):
CAPACITOR_DATA = []
REACTOR_DATA = []
SERIES_COMP_DATA = []
SERIES_REACTOR_DATA = []
SHUNT_DATA = []''')
    
    doc.add_heading('2.2 Bus Type Codes', level=2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    types = [('Type', 'Description'), ('1', 'PQ Bus (Load bus) - P and Q specified'), 
             ('2', 'PV Bus (Generator) - P and V specified'), ('3', 'Slack Bus - V and angle specified')]
    for i, (t, d) in enumerate(types):
        table.rows[i].cells[0].text = t
        table.rows[i].cells[1].text = d
        if i == 0:
            set_cell_shading(table.rows[i].cells[0], '2E75B6')
            set_cell_shading(table.rows[i].cells[1], '2E75B6')
    
    doc.add_heading('2.3 Input Value Units', level=2)
    
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    units = [
        ('Parameter', 'Unit', 'Notes'),
        ('Voltage (V_init)', 'per-unit', 'Typically 0.95-1.05'),
        ('Angle', 'degrees', 'Slack bus = 0°'),
        ('Impedance (R, X)', 'per-unit', 'On system base'),
        ('Susceptance (B)', 'per-unit', 'Total line charging'),
        ('Power (P, Q)', 'MW, Mvar', 'Actual values'),
        ('Length', 'km', 'For per-km calculations'),
        ('Rating (rateA)', 'MVA', 'Thermal limit')
    ]
    for i, row_data in enumerate(units):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                set_cell_shading(table.rows[i].cells[j], '2E75B6')
    
    doc.add_page_break()
    
    # SECTION 3: RUNNING THE SYSTEM
    doc.add_heading('3. RUNNING THE SYSTEM', level=1)
    
    doc.add_heading('3.1 Prerequisites', level=2)
    add_code_block(doc, '''Required Python Packages:
    pip install numpy

Optional (for PYPOWER conversion):
    pip install pypower''')
    
    doc.add_heading('3.2 Interactive Mode', level=2)
    add_code_block(doc, '''Step 1: Open command prompt/terminal
Step 2: Navigate to the folder containing the scripts
Step 3: Run:
    python run_power_flow.py

Step 4: Enter requested information:
    - Input file path (e.g., "E:\\data\\case300.py")
    - Variation strength (1=5%, 2=10%, 3=15%)
    - Number of samples (default: 200)

Step 5: Wait for completion
Step 6: Check output folder for results''')
    
    doc.add_heading('3.3 Command Line Mode', level=2)
    add_code_block(doc, '''Basic Usage:
    python run_power_flow.py -i <input_file> -v <variation> -n <samples>

Arguments:
    -i, --input      Path to input Python file
    -v, --variation  Variation strength in percent (5, 10, or 15)
    -n, --samples    Number of samples to generate

Examples:
    # Full parameters
    python run_power_flow.py -i "E:\\data\\case300.py" -v 10 -n 500

    # Default variation (10%), custom samples
    python run_power_flow.py -i case300.py -n 100

    # Short path
    python run_power_flow.py -i my_system.py -v 15''')
    
    doc.add_heading('3.4 Variation Strength Explained', level=2)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    var_data = [
        ('Option', 'Strength', 'Voltage Range', 'Angle Range'),
        ('1 (5%)', 'Small', '0.95 - 1.05 pu', '±1.5°'),
        ('2 (10%)', 'Medium', '0.90 - 1.10 pu', '±3.0°'),
        ('3 (15%)', 'Large', '0.85 - 1.15 pu', '±4.5°')
    ]
    for i, row in enumerate(var_data):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            if i == 0:
                set_cell_shading(table.rows[i].cells[j], '2E75B6')
    
    doc.add_paragraph()
    doc.add_paragraph('Higher variation produces more diverse samples but may include unrealistic operating points.')
    
    doc.add_page_break()
    
    # SECTION 4: OUTPUTS
    doc.add_heading('4. OUTPUT SPECIFICATIONS', level=1)
    
    doc.add_heading('4.1 Output Files Generated', level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    outputs = [
        ('File', 'Format', 'Contents'),
        ('*_ALL_DATA.csv', 'CSV', 'Complete input+output for all components'),
        ('*_ALL_DATA.py', 'Python', 'Same data as CSV, loadable as module'),
        ('*_OUTPUTS_ONLY.py', 'Python', 'Only calculation results (compact)'),
        ('*_SUMMARY.txt', 'Text', 'System summary and key metrics')
    ]
    for i, row in enumerate(outputs):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            if i == 0:
                set_cell_shading(table.rows[i].cells[j], '2E75B6')
    
    doc.add_heading('4.2 Output Folder Structure', level=2)
    add_code_block(doc, '''After running, outputs are saved in:
    <input_file_directory>/<input_filename>/
        ├── case300_var10_20240115_143000_ALL_DATA.csv
        ├── case300_var10_20240115_143000_ALL_DATA.py
        ├── case300_var10_20240115_143000_OUTPUTS_ONLY.py
        └── case300_var10_20240115_143000_SUMMARY.txt

Note: Folder is named after the input file.''')
    
    doc.add_heading('4.3 Output Data Fields', level=2)
    
    doc.add_paragraph('BUS Outputs:', style='List Bullet')
    add_code_block(doc, '''O_V_final_pu      - Final voltage magnitude (per-unit)
O_V_final_kV      - Final voltage magnitude (kV)
O_angle_deg       - Voltage angle (degrees)
O_voltage_status  - LOW/NORMAL/HIGH (based on 0.95-1.05 pu range)
O_angle_status    - NORMAL if |angle| < 30°, else HIGH''')
    
    doc.add_paragraph('BRANCH Outputs:', style='List Bullet')
    add_code_block(doc, '''O_P_fwd_MW        - Active power from→to (MW)
O_Q_fwd_Mvar      - Reactive power from→to (Mvar)
O_MVA_fwd_MVA     - Apparent power from→to (MVA)
O_P_rev_MW        - Active power to→from (MW)
O_Q_rev_Mvar      - Reactive power to→from (Mvar)
O_MVA_rev_MVA     - Apparent power to→from (MVA)
O_losses_MW       - Total losses (MW)
O_loading_pct     - Loading percentage
O_status          - NORMAL/MODERATE/HIGH/OVERLOADED''')
    
    doc.add_paragraph('TRANSFORMER Additional Outputs:', style='List Bullet')
    add_code_block(doc, '''O_tap_position    - Current tap ratio value
O_tap_step        - Integer tap step number
O_losses_MVA      - Losses in MVA (not just MW)''')
    
    doc.add_page_break()
    
    # SECTION 5: CONVERTING PYPOWER FILES
    doc.add_heading('5. CONVERTING PYPOWER FILES', level=1)
    
    doc.add_heading('5.1 Single File Conversion', level=2)
    add_code_block(doc, '''Run the converter:
    python frompypowerformatetothis.py

Select option 1 (Single File)
Enter the path to your PYPOWER .py file

Output: converted_<original_filename>.py in same folder''')
    
    doc.add_heading('5.2 Batch Conversion', level=2)
    add_code_block(doc, '''Run the converter:
    python frompypowerformatetothis.py

Select option 2 (Batch Mode)
Enter the folder path containing PYPOWER files

All .py files will be converted with "converted_" prefix''')
    
    doc.add_heading('5.3 PYPOWER to Custom Format Mapping', level=2)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    mapping = [
        ('PYPOWER Field', 'Index', 'Custom Format'),
        ('bus[:,0]', '0', 'bus_num'),
        ('bus[:,1]', '1', 'bus_type'),
        ('bus[:,2]', '2', 'P_demand → LOAD_DATA'),
        ('bus[:,3]', '3', 'Q_demand → LOAD_DATA'),
        ('bus[:,4]', '4', 'shunt_G'),
        ('bus[:,5]', '5', 'shunt_B → CAPACITOR/REACTOR')
    ]
    for i, row in enumerate(mapping):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            if i == 0:
                set_cell_shading(table.rows[i].cells[j], '2E75B6')
    
    doc.add_page_break()
    
    # SECTION 6: TROUBLESHOOTING
    doc.add_heading('6. TROUBLESHOOTING', level=1)
    
    problems = [
        ('Error: File not found', 'Check file path. Use forward slashes or escape backslashes: "E:\\\\data\\\\file.py"'),
        ('Error: Could not find valid PYPOWER dictionary', 'The input file must have a function that returns a ppc dictionary'),
        ('No valid samples generated', 'Reduce variation strength. Check input data for invalid impedances (R=0, X=0)'),
        ('NaN values in output', 'Check for zero-impedance branches or disconnected buses'),
        ('Loading > 100% everywhere', 'Check if rateA values are correct and not zero'),
        ('Voltage values seem wrong', 'Verify base_kV values match actual system voltages'),
    ]
    
    for problem, solution in problems:
        p = doc.add_paragraph()
        p.add_run(f'Problem: {problem}').bold = True
        doc.add_paragraph(f'Solution: {solution}')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # SECTION 7: FAQ
    doc.add_heading('7. FREQUENTLY ASKED QUESTIONS', level=1)
    
    faqs = [
        ('Q: What is "inverse" power flow?', 
         'A: Traditional power flow solves for voltages given powers. Inverse power flow calculates powers given voltages. This is useful for generating training data for machine learning models.'),
        
        ('Q: Why do I need multiple samples?',
         'A: Each sample represents a different operating point. More samples provide better statistical coverage of the operating envelope.'),
        
        ('Q: Can I add my own components?',
         'A: Yes. Modify input_reader.py to parse new data types and report_writer.py to output them. Do NOT modify power_flow_engine.py.'),
        
        ('Q: What happens if rateA is 0?',
         'A: Loading percentage is set to 0 (no limit). Losses are still calculated correctly.'),
        
        ('Q: How accurate are the calculations?',
         'A: The calculations use standard AC power flow equations with the π-model for lines and ideal transformer model. Accuracy depends on input data quality.'),
        
        ('Q: Can I use this for real-time applications?',
         'A: The current implementation is for offline analysis. For real-time use, you would need to integrate it with SCADA data and add convergence checking.'),
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        doc.add_paragraph(a)
        doc.add_paragraph()
    
    doc.save('02_User_Manual.docx')
    print("Created: 02_User_Manual.docx")


def create_simplified_document():
    """DOCUMENT 3: Simplified Summary (2-3 pages)"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Title
    title = doc.add_heading('POWER FLOW SYSTEM - QUICK REFERENCE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('All 5 Files Summary in 3 Pages').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # FILE 1
    doc.add_heading('FILE 1: frompypowerformatetothis.py', level=1)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Converts PYPOWER format files to custom input format')
    
    p = doc.add_paragraph()
    p.add_run('Key Functions:').bold = True
    add_code_block(doc, '''convert_pypower_to_custom(file_path)
  → Reads PYPOWER ppc dictionary
  → Extracts bus, gen, branch data
  → Separates lines (ratio=0) from transformers (ratio≠0)
  → Converts shunt_B to capacitor/reactor data
  → Writes converted_<filename>.py''')
    
    p = doc.add_paragraph()
    p.add_run('Usage: ').bold = True
    p.add_run('Run script → Choose 1 (single) or 2 (batch) → Provide path')
    
    doc.add_paragraph()
    
    # FILE 2
    doc.add_heading('FILE 2: input_reader.py', level=1)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Parses custom format input files and prepares data for engine')
    
    p = doc.add_paragraph()
    p.add_run('Key Functions:').bold = True
    add_code_block(doc, '''read_from_python(file_path)
  → Loads .py file as module
  → Extracts all DATA variables into dictionaries
  → Returns structured data dict

convert_to_engine_format(input_data)
  → Converts to simple format for engine
  → Calculates line totals: R×length, X×length, B×length
  → Returns bus_data dict and branch_data list''')
    
    p = doc.add_paragraph()
    p.add_run('Input Format Example:').bold = True
    add_code_block(doc, '''BUS_DATA = [
    [1, "Bus_1", 3, 132, 1.05, 0.0, 0, 0, 1, 1, 1],
    [2, "Bus_2", 1, 132, 1.00, 0.0, 0, 0, 1, 1, 1]
]
LINE_DATA = [
    [1, "Line_1_2", 1, 2, 100, 0.0001, 0.0003, 0.00001, 200, 1, 1, 1, 1]
]''')
    
    doc.add_paragraph()
    
    # FILE 3
    doc.add_heading('FILE 3: power_flow_engine.py', level=1)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Core calculation engine - DO NOT MODIFY')
    
    p = doc.add_paragraph()
    p.add_run('Key Calculations:').bold = True
    add_code_block(doc, '''build_admittance_matrix()
  Ybus[i][i] += y + jB          (line diagonal)
  Ybus[i][j] -= y               (line off-diagonal)
  Ybus[i][i] += y/a²           (transformer diagonal)
  Ybus[i][j] -= y/a            (transformer off-diagonal)

calculate_power(V_mag, V_angle)
  S = V × (Ybus × V)*
  P = real(S) × baseMVA
  Q = imag(S) × baseMVA

calculate_bidirectional_flows()
  Line:      S_fwd = Vi × [(Vi-Vj)×y + Vi×jB/2]*
  Xfmr:      S_fwd = Vi × [(Vi/a - Vj)×y]*
  Losses = S_fwd + S_rev''')
    
    p = doc.add_paragraph()
    p.add_run('Sample Generation:').bold = True
    add_code_block(doc, '''generate_sample(variation_strength)
  → Applies random ±variation to V_mag and V_angle
  → Clips voltage to 0.90-1.10 pu
  → Returns power injections and branch flows''')
    
    doc.add_page_break()
    
    # FILE 4
    doc.add_heading('FILE 4: report_writer.py', level=1)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Generates output reports in multiple formats')
    
    p = doc.add_paragraph()
    p.add_run('Output Methods:').bold = True
    add_code_block(doc, '''write_csv_all_data()      → *_ALL_DATA.csv
write_py_all_data()       → *_ALL_DATA.py  
write_py_output_only()    → *_OUTPUTS_ONLY.py
write_txt_summary()       → *_SUMMARY.txt''')
    
    p = doc.add_paragraph()
    p.add_run('Output Includes:').bold = True
    add_code_block(doc, '''• Bus: V_pu, V_kV, angle, status (LOW/NORMAL/HIGH)
• Generator: P_out, Q_out, V_actual, Q_status
• Load: P_supplied, Q_supplied, supply_status
• Line: P_fwd, Q_fwd, MVA_fwd, P_rev, Q_rev, MVA_rev, 
        losses, loss_per_km, loading_%
• Transformer: Same as line + tap_ratio, tap_step, loss_MVA
• Capacitor/Reactor: Q_injected/absorbed, V_actual
• System Summary: Total gen, load, losses, loss_%, voltages''')
    
    doc.add_paragraph()
    
    # FILE 5
    doc.add_heading('FILE 5: run_power_flow.py', level=1)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Main orchestrator - connects all components')
    
    p = doc.add_paragraph()
    p.add_run('Execution Modes:').bold = True
    add_code_block(doc, '''Interactive:
  python run_power_flow.py
  → Prompts for file, variation, samples

Command Line:
  python run_power_flow.py -i file.py -v 10 -n 200
  -i: input file path
  -v: variation (5, 10, or 15 percent)
  -n: number of samples''')
    
    p = doc.add_paragraph()
    p.add_run('Workflow:').bold = True
    add_code_block(doc, '''1. Read input file (input_reader.py)
2. Convert to engine format (input_reader.py)
3. Initialize engine (power_flow_engine.py)
4. Run batch simulation (power_flow_engine.py)
5. Generate reports (report_writer.py)
6. Save to <input_filename>/ folder''')
    
    doc.add_paragraph()
    
    # QUICK REFERENCE TABLE
    doc.add_heading('QUICK REFERENCE: DATA FLOW', level=1)
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    flow_data = [
        ('Stage', 'Input', 'Process', 'Output'),
        ('1. Convert', 'PYPOWER .py', 'frompypowerformatetothis.py', 'Custom .py'),
        ('2. Read', 'Custom .py', 'input_reader.py', 'Data dict'),
        ('3. Calculate', 'Data dict', 'power_flow_engine.py', 'Samples'),
        ('4. Report', 'Samples', 'report_writer.py', 'CSV/PY/TXT'),
        ('5. Run All', 'Command', 'run_power_flow.py', 'All files')
    ]
    for i, row in enumerate(flow_data):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            if i == 0:
                set_cell_shading(table.rows[i].cells[j], '2E75B6')
    
    doc.add_paragraph()
    
    # KEY FORMULAS
    doc.add_heading('KEY FORMULAS AT A GLANCE', level=1)
    add_code_block(doc, '''Admittance:        y = 1/(R + jX) = (R - jX)/(R² + X²)
Power Injection:  S = V × (Ybus × V)*
Line Forward:     S_fwd = Vi × [(Vi-Vj)×y + Vi×jB/2]*
Xfmr Forward:     S_fwd = Vi × [(Vi/a - Vj)×y]*
Losses:           Loss = S_fwd + S_rev
Loading:          Loading_% = |S_fwd| × baseMVA / rateA × 100
Capacitor Q:      Q_out = Q_rated × V²
Voltage kV:       V_kV = V_pu × base_kV''')
    
    doc.save('03_Simplified_Summary.docx')
    print("Created: 03_Simplified_Summary.docx")


if __name__ == "__main__":
    print("="*60)
    print("GENERATING POWER FLOW DOCUMENTATION")
    print("="*60)
    print()
    
    print("Creating Document 1: Main Calculations...")
    create_calculation_document()
    
    print("Creating Document 2: User Manual...")
    create_user_manual_document()
    
    print("Creating Document 3: Simplified Summary...")
    create_simplified_document()
    
    print()
    print("="*60)
    print("ALL DOCUMENTS GENERATED SUCCESSFULLY!")
    print("="*60)
    print()
    print("Files created:")
    print("  📄 01_Main_Calculations.docx  (Detailed math & code)")
    print("  📄 02_User_Manual.docx        (How to use)")
    print("  📄 03_Simplified_Summary.docx (2-3 page quick ref)")